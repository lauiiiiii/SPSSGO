# -*- coding: utf-8 -*-
import os
import tempfile
import unittest

from docx import Document

from backend.word_report import (
    BLACK,
    BODY_FONT_SIZE,
    HEADING_1_FONT_SIZE,
    HEADING_2_FONT_SIZE,
    TABLE_FONT_SIZE,
    TITLE_FONT_SIZE,
    generate_report,
)


class WordReportStyleTests(unittest.TestCase):
    def test_generate_report_uses_black_academic_fonts(self):
        results = [
            {
                "name": "正态性分析",
                "sections": [
                    {
                        "type": "table",
                        "title": "正态性检验结果",
                        "headers": ["变量", "N"],
                        "rows": [["q1_1", 10]],
                        "description": "这是一段结果描述。",
                    },
                    {
                        "type": "advice",
                        "title": "分析建议",
                        "content": "这是一段正文内容。",
                    },
                ],
            }
        ]

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name
        try:
            generate_report(results, output_path, title="数据分析结果")
            doc = Document(output_path)

            self.assertEqual(doc.paragraphs[0].text, "数据分析结果")
            self.assertEqual(doc.paragraphs[0].style.name, "Normal")
            self.assertEqual(doc.paragraphs[0].runs[0].font.color.rgb, BLACK)
            self.assertEqual(doc.paragraphs[0].runs[0].font.size, TITLE_FONT_SIZE)

            self.assertEqual(doc.paragraphs[1].text, "正态性分析")
            self.assertEqual(doc.paragraphs[1].style.name, "Normal")
            self.assertEqual(doc.paragraphs[1].runs[0].font.color.rgb, BLACK)
            self.assertEqual(doc.paragraphs[1].runs[0].font.size, HEADING_1_FONT_SIZE)

            self.assertEqual(doc.paragraphs[2].text, "正态性检验结果")
            self.assertEqual(doc.paragraphs[2].style.name, "Normal")
            self.assertEqual(doc.paragraphs[2].runs[0].font.color.rgb, BLACK)
            self.assertEqual(doc.paragraphs[2].runs[0].font.size, HEADING_2_FONT_SIZE)

            table_header_run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
            self.assertEqual(table_header_run.font.color.rgb, BLACK)
            self.assertEqual(table_header_run.font.size, TABLE_FONT_SIZE)

            description_paragraph = next(item for item in doc.paragraphs if item.text == "这是一段结果描述。")
            self.assertEqual(description_paragraph.runs[0].font.color.rgb, BLACK)
            self.assertEqual(description_paragraph.runs[0].font.size, BODY_FONT_SIZE)
        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
