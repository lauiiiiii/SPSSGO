# -*- coding: utf-8 -*-
"""
Word 三线表报告生成
"""
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm

from backend.domain import normalize_analysis_sections

BODY_FONT_SIZE = Pt(12)       # 小四
TABLE_FONT_SIZE = Pt(10.5)    # 五号
NOTE_FONT_SIZE = Pt(9)        # 小五
TITLE_FONT_SIZE = Pt(22)      # 二号
HEADING_1_FONT_SIZE = Pt(16)  # 三号
HEADING_2_FONT_SIZE = Pt(14)  # 四号
BLACK = RGBColor(0, 0, 0)


def _style_run(run, *, east_asia_font: str, western_font: str = "Times New Roman", size=None, bold: bool = False):
    """统一字体和颜色，别继续吃 Word 默认蓝色标题样式。"""
    run.bold = bold
    run.font.name = western_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = size


def _add_heading(doc: Document, text: str, *, size, center: bool = False):
    """标题直接写普通段落，别再用 add_heading 吃模板自带颜色。"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    _style_run(run, east_asia_font="黑体", size=size, bold=True)
    return paragraph


def _add_body_paragraph(doc: Document, text: str, *, first_line_indent: bool = True, size=BODY_FONT_SIZE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    run = paragraph.add_run(text)
    _style_run(run, east_asia_font="宋体", size=size)
    return paragraph


def _cell_text(cell):
    if isinstance(cell, dict):
        return str(cell.get('text', '') or '')
    return str(cell) if cell is not None else ''


def _cell_span(cell, key):
    if isinstance(cell, dict):
        try:
            return max(1, int(cell.get(key, 1) or 1))
        except (TypeError, ValueError):
            return 1
    return 1


def _apply_table_outer_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '</w:tblBorders>' % nsdecls("w")
    ))


def _apply_cell_bottom_border(tc, weight=6):
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcPr.append(parse_xml(
        '<w:tcBorders %s>'
        '  <w:bottom w:val="single" w:sz="%d" w:space="0" w:color="000000"/>'
        '</w:tcBorders>' % (nsdecls("w"), weight)
    ))


def _make_three_line_table(doc, headers, data_rows):
    """创建标准 APA 三线表：顶线(粗) - 表头底线(细) - 底线(粗)，无竖线"""
    return _make_three_line_table_ex(doc, headers, data_rows, [])


def _make_three_line_table_ex(doc, headers, data_rows, header_rows=None):
    """创建支持 headerRows（含 colspan 合并）的标准三线表。"""
    header_rows = header_rows or []
    n_cols = len(headers) if headers else max(
        (sum(_cell_span(c, 'colspan') for c in hr) for hr in header_rows if hr),
        default=max((len(r) for r in data_rows if r), default=1)
    )
    if n_cols < 1:
        return None

    n_head = len(header_rows) if header_rows else 1
    n_data = len(data_rows)
    table = doc.add_table(rows=n_head + n_data, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # --- 表头行 ---
    if header_rows:
        for ri, hr in enumerate(header_rows):
            col_cursor = 0
            for cell_def in hr:
                if col_cursor >= n_cols:
                    break
                cs = _cell_span(cell_def, 'colspan')
                rs = _cell_span(cell_def, 'rowspan')
                cell = table.rows[ri].cells[col_cursor]
                if cs > 1 and col_cursor + cs - 1 < n_cols:
                    cell = cell.merge(table.rows[ri].cells[col_cursor + cs - 1])
                if rs > 1 and ri + rs - 1 < n_head:
                    cell = cell.merge(table.rows[ri + rs - 1].cells[col_cursor])
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(_cell_text(cell_def))
                _style_run(run, east_asia_font="宋体", size=TABLE_FONT_SIZE, bold=True)
                col_cursor += cs
    else:
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(h))
            _style_run(run, east_asia_font="宋体", size=TABLE_FONT_SIZE, bold=True)

    # --- 数据行 ---
    for i, row_data in enumerate(data_rows):
        for j in range(min(len(row_data), n_cols)):
            cell = table.rows[n_head + i].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_cell_text(row_data[j]))
            _style_run(run, east_asia_font="宋体", size=TABLE_FONT_SIZE)

    _apply_table_outer_borders(table)

    # --- 最后一个表头行底部细线 ---
    last_head_row = table.rows[n_head - 1]
    for j in range(n_cols):
        _apply_cell_bottom_border(last_head_row.cells[j]._tc)

    return table


def generate_report(results: list[dict], output_path: str, title: str = "数据分析结果"):
    doc = Document()
    _configure_academic_page(doc)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = BODY_FONT_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.color.rgb = BLACK

    _add_heading(doc, title, size=TITLE_FONT_SIZE, center=True)

    for idx, r in enumerate(results):
        name = r.get("name") or r.get("analysis_name") or f"分析 {idx+1}"
        _add_heading(doc, name, size=HEADING_1_FONT_SIZE)

        sections = normalize_analysis_sections(r.get("sections")) or []
        if sections:
            for sec in sections:
                sec_type = sec.get("type")
                sec_title = sec.get("title", "") or ""

                if sec_type == "table":
                    inline_title = sec.get("inlineTitle") or False
                    if sec_title and not inline_title:
                        _add_heading(doc, sec_title, size=HEADING_2_FONT_SIZE)
                    headers = sec.get("headers") or []
                    rows = sec.get("exportRows") or sec.get("rows") or []
                    header_rows = sec.get("headerRows") or []
                    if rows and (headers or header_rows):
                        _make_three_line_table_ex(doc, headers, rows, header_rows)
                    note = sec.get("note", "") or ""
                    if note:
                        _add_body_paragraph(doc, note, first_line_indent=False, size=NOTE_FONT_SIZE)
                    desc = sec.get("description", "") or ""
                    if desc:
                        _add_body_paragraph(doc, desc)
                elif sec_type in {"advice", "smart_analysis"}:
                    if sec_title:
                        _add_heading(doc, sec_title, size=HEADING_2_FONT_SIZE)
                    content = sec.get("content", "") or ""
                    if content:
                        _add_body_paragraph(doc, content)
                elif sec_type == "references":
                    if sec_title:
                        _add_heading(doc, sec_title, size=HEADING_2_FONT_SIZE)
                    items = sec.get("items") or []
                    for item in items:
                        _add_body_paragraph(doc, f"{item}", first_line_indent=False, size=NOTE_FONT_SIZE)
        else:
            headers = r.get("headers") or r.get("table_headers") or []
            rows = r.get("rows") or r.get("table_rows") or []
            if headers and rows:
                _make_three_line_table(doc, headers, rows)

            desc = r.get("description", "") or ""
            if desc:
                _add_body_paragraph(doc, desc)

    doc.save(output_path)
    return output_path


def _configure_academic_page(doc: Document):
    """统一成 A4 论文页式，别让导出的文档继续走默认办公模板。"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

